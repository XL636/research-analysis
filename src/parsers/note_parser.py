"""笔记解析器 - 解析 Markdown/TXT/DOCX 文件."""

from __future__ import annotations

import re
from pathlib import Path

from loguru import logger

from src.core.models import DocumentSection, FileType, ParsedDocument


def parse_note(file_path: str) -> ParsedDocument:
    """解析笔记文件（Markdown/TXT/DOCX）."""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    suffix = path.suffix.lower()
    logger.info(f"Parsing note [{suffix}]: {path.name}")

    if suffix == ".docx":
        return _parse_docx(file_path)
    elif suffix in (".md", ".markdown"):
        return _parse_markdown(file_path)
    else:  # .txt and others
        return _parse_txt(file_path)


def _parse_markdown(file_path: str) -> ParsedDocument:
    """解析 Markdown 文件."""
    path = Path(file_path)
    text = path.read_text(encoding="utf-8")

    title = ""
    sections: list[DocumentSection] = []

    # 提取标题和章节
    lines = text.split("\n")
    current_title = ""
    current_lines: list[str] = []
    current_level = 1

    for line in lines:
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            # 保存前一个 section
            if current_lines or current_title:
                sections.append(DocumentSection(
                    title=current_title,
                    content="\n".join(current_lines).strip(),
                    level=current_level,
                ))
                current_lines = []

            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()

            if not title and level == 1:
                title = heading_text

            current_title = heading_text
            current_level = level
        else:
            current_lines.append(line)

    # 最后一个 section
    if current_lines or current_title:
        sections.append(DocumentSection(
            title=current_title,
            content="\n".join(current_lines).strip(),
            level=current_level,
        ))

    return ParsedDocument(
        file_path=file_path,
        file_type=FileType.MARKDOWN,
        title=title or path.stem,
        sections=sections,
        raw_text=text,
    )


def _parse_txt(file_path: str) -> ParsedDocument:
    """解析纯文本文件."""
    path = Path(file_path)
    text = path.read_text(encoding="utf-8")

    # 尝试用空行分段
    paragraphs = re.split(r"\n\s*\n", text)
    sections = []
    for i, para in enumerate(paragraphs):
        para = para.strip()
        if para:
            sections.append(DocumentSection(
                title=f"段落 {i + 1}" if len(paragraphs) > 1 else "",
                content=para,
                level=2,
            ))

    return ParsedDocument(
        file_path=file_path,
        file_type=FileType.TXT,
        title=path.stem,
        sections=sections,
        raw_text=text,
    )


def _parse_docx(file_path: str) -> ParsedDocument:
    """解析 DOCX 文件."""
    try:
        import docx
    except ImportError:
        logger.warning("python-docx not installed, reading as plain text fallback")
        return _parse_txt(file_path)

    path = Path(file_path)
    doc = docx.Document(file_path)

    title = ""
    sections: list[DocumentSection] = []
    all_text_parts: list[str] = []
    current_title = ""
    current_lines: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        all_text_parts.append(text)

        # 检测标题样式
        style_name = para.style.name if para.style else ""
        is_heading = "Heading" in style_name or "heading" in style_name

        if is_heading:
            if current_lines or current_title:
                sections.append(DocumentSection(
                    title=current_title,
                    content="\n".join(current_lines),
                    level=2,
                ))
                current_lines = []

            if not title and ("Heading 1" in style_name or "Title" in style_name):
                title = text
            current_title = text
        else:
            current_lines.append(text)

    if current_lines or current_title:
        sections.append(DocumentSection(
            title=current_title,
            content="\n".join(current_lines),
            level=2,
        ))

    raw_text = "\n\n".join(all_text_parts)

    # 元数据
    metadata: dict[str, str] = {}
    authors: list[str] = []
    if doc.core_properties:
        if doc.core_properties.title:
            title = title or doc.core_properties.title
            metadata["docx_title"] = doc.core_properties.title
        if doc.core_properties.author:
            authors = [doc.core_properties.author]

    return ParsedDocument(
        file_path=file_path,
        file_type=FileType.DOCX,
        title=title or path.stem,
        authors=authors,
        sections=sections,
        raw_text=raw_text,
        metadata=metadata,
    )
