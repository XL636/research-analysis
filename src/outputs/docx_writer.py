"""DOCX 输出适配器 - 生成 Word 文档."""

from __future__ import annotations

import re
from pathlib import Path

from loguru import logger

from src.core.models import Report


def write_docx(report: Report, output_path: str) -> str:
    """将报告写入 DOCX 文件.

    Args:
        report: 报告对象
        output_path: 输出文件路径

    Returns:
        实际写入的文件路径
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for DOCX output. Install with: uv add python-docx")

    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # 标题
    doc.add_heading(report.title, level=0)

    # 元信息
    meta = f"生成时间：{report.generated_at.strftime('%Y-%m-%d %H:%M')}"
    if report.source_files:
        meta += f"\n来源文件：{', '.join(report.source_files)}"
    doc.add_paragraph(meta, style="Subtitle")

    # 解析 Markdown 内容并转为 DOCX
    _markdown_to_docx(doc, report.content)

    doc.save(output_path)
    logger.info(f"DOCX report written to: {path}")
    return str(path)


def _markdown_to_docx(doc, markdown_text: str) -> None:
    """简易 Markdown → DOCX 转换."""
    lines = markdown_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # 标题
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            doc.add_heading(heading_match.group(2), level=min(level, 4))
            i += 1
            continue

        # 无序列表
        list_match = re.match(r"^\s*[-*]\s+(.+)$", line)
        if list_match:
            doc.add_paragraph(list_match.group(1), style="List Bullet")
            i += 1
            continue

        # 有序列表
        olist_match = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if olist_match:
            doc.add_paragraph(olist_match.group(1), style="List Number")
            i += 1
            continue

        # 分隔线
        if re.match(r"^---+$", line.strip()):
            doc.add_paragraph("─" * 40)
            i += 1
            continue

        # 引用块
        quote_match = re.match(r"^>\s*(.*)$", line)
        if quote_match:
            doc.add_paragraph(quote_match.group(1), style="Quote")
            i += 1
            continue

        # 普通段落
        stripped = line.strip()
        if stripped:
            # 清理 Markdown 加粗/斜体标记
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            text = re.sub(r"\*(.+?)\*", r"\1", text)
            text = re.sub(r"`(.+?)`", r"\1", text)
            doc.add_paragraph(text)

        i += 1
